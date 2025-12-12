import uvicorn
import fastapi
import time
import re
import io
import os
import glob
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
import logging

# basic logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
import docx
import PyPDF2
from typing import List, Dict, Any, Set, Optional, Tuple
from pydantic import BaseModel

# fastapi app setup
app = FastAPI(
    title="Intelligent CV Analyzer (v2)",
    description="Batch-analyzes CVs and provides a ranked list."
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# store cvs in memory
cv_database: Dict[str, str] = {}

# job descriptions with keywords for matching
job_descriptions = [
    {
        "id": "jd1",
        "title": "AI/ML Engineer",
        "keywords": [
            "python", "machine learning", "deep learning", "tensorflow", "pytorch",
            "scikit-learn", "pandas", "numpy", "opencv", "computer vision", "nlp",
            "langchain", "fastapi", "docker", "git"
        ]
    },
    {
        "id": "jd2",
        "title": "Data Scientist",
        "keywords": [
            "python", "r", "sql", "data analysis", "pandas", "numpy", "matplotlib",
            "seaborn", "scikit-learn", "data mining", "power bi", "tableau",
            "statistics", "regression", "classification", "clustering"
        ]
    },
    {
        "id": "jd3",
        "title": "Full-Stack Web Developer",
        "keywords": [
            "html", "css", "javascript", "react", "node.js", "python", "django",
            "flask", "fastapi", "mongodb", "mysql", "postgresql", "rest apis",
            "git", "docker", "figma"
        ]
    }
]

# file validation functions

def is_valid_cv_filename(filename: str) -> bool:
    # checks if filename follows pattern: 2digits + lowercase letter + 4digits
    # check if original has uppercase letters first
    original_base = filename
    original_base = re.sub(r'\.(pdf|docx?)$', '', original_base)
    original_base = re.sub(r'[_\-\s]*cv[_\-\s]*', '', original_base, flags=re.IGNORECASE)
    original_base = re.sub(r'\([0-9]+\)$', '', original_base)
    original_base = re.sub(r'[_\-\s]+$', '', original_base)
    
    # reject if uppercase letters found
    if re.search(r'[A-Z]', original_base):
        return False
    
    # clean up filename for validation
    base_name = filename.lower()
    base_name = re.sub(r'\.(pdf|docx?)$', '', base_name)
    base_name = re.sub(r'[_\-\s]*cv[_\-\s]*', '', base_name)
    base_name = re.sub(r'\([0-9]+\)$', '', base_name)
    base_name = re.sub(r'[_\-\s]+$', '', base_name)
    
    # check pattern: 2 digits + lowercase letter + 4 digits
    pattern = r'^[0-9]{2}[a-z][0-9]{4}$'
    return bool(re.match(pattern, base_name))

def extract_student_id(filename: str) -> Optional[str]:
    base_name = filename.lower()
    base_name = re.sub(r'\.(pdf|docx?)$', '', base_name)
    base_name = re.sub(r'[_\-\s]*cv[_\-\s]*', '', base_name)
    base_name = re.sub(r'\([0-9]+\)$', '', base_name)
    base_name = re.sub(r'[_\-\s]+$', '', base_name)
    
    match = re.match(r'^([0-9]{2}[a-z][0-9]{4})', base_name)
    return match.group(1) if match else None

def get_submission_number(filename: str) -> int:
    match = re.search(r'\(([0-9]+)\)', filename)
    return int(match.group(1)) if match else 0

def find_latest_submission(filenames: List[str]) -> Dict[str, str]:
    student_files = {}
    
    for filename in filenames:
        if not is_valid_cv_filename(filename):
            continue
            
        student_id = extract_student_id(filename)
        if not student_id:
            continue
            
        submission_num = get_submission_number(filename)
        is_pdf = filename.lower().endswith('.pdf')
        
        if student_id not in student_files:
            student_files[student_id] = {
                'filename': filename,
                'submission_num': submission_num,
                'is_pdf': is_pdf
            }
        else:
            current = student_files[student_id]
            
            # keep latest submission
            if submission_num > current['submission_num']:
                student_files[student_id] = {
                    'filename': filename,
                    'submission_num': submission_num,
                    'is_pdf': is_pdf
                }
            # if same number, prefer pdf
            elif submission_num == current['submission_num'] and is_pdf and not current['is_pdf']:
                student_files[student_id] = {
                    'filename': filename,
                    'submission_num': submission_num,
                    'is_pdf': is_pdf
                }
    
    # return just filenames
    return {sid: data['filename'] for sid, data in student_files.items()}

def load_dataset_files(dataset_path: str = "./DataSet") -> Tuple[List[str], List[str], List[str]]:
    if not os.path.exists(dataset_path):
        return [], [], []
    
    # get all cv files
    pdf_files = glob.glob(os.path.join(dataset_path, "*.pdf"))
    docx_files = glob.glob(os.path.join(dataset_path, "*.docx"))
    doc_files = glob.glob(os.path.join(dataset_path, "*.doc"))
    
    all_files = pdf_files + docx_files + doc_files
    filenames = [os.path.basename(f) for f in all_files]
    
    # separate valid and invalid files
    valid_files = []
    rejected_files = []
    
    for filename in filenames:
        if is_valid_cv_filename(filename):
            valid_files.append(filename)
        else:
            rejected_files.append(filename)
    
    # keep only latest submissions
    latest_submissions = find_latest_submission(valid_files)
    filtered_files = list(latest_submissions.values())
    
    return valid_files, filtered_files, rejected_files

# text processing functions

def normalize_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^\w\s]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

async def extract_text(file: UploadFile) -> Tuple[str, Optional[str]]:
    filename = file.filename
    content = await file.read()
    
    try:
        if filename.lower().endswith((".docx", ".doc")):
            try:
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if not text.strip():
                    return "", f"Document appears to be empty: {filename}"
                return text, None
            except Exception as e:
                return "", f"Failed to parse Word document {filename}: {str(e)}"
                
        elif filename.lower().endswith(".pdf"):
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                
                if len(pdf_reader.pages) == 0:
                    return "", f"PDF has no pages: {filename}"
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"page {page_num + 1} extraction failed in {filename}")
                        continue
                
                if not text.strip():
                    return "", f"No text could be extracted from PDF: {filename}"
                    
                return text, None
            except Exception as e:
                return "", f"Failed to parse PDF {filename}: {str(e)}"
        else:
            return "", f"Unsupported file format: {filename}"
            
    except Exception as e:
        return "", f"Unexpected error processing {filename}: {str(e)}"

def extract_text_from_file(file_path: str) -> Tuple[str, Optional[str]]:
    if not os.path.exists(file_path):
        return "", f"File not found: {file_path}"
    
    try:
        filename = os.path.basename(file_path)
        
        if filename.lower().endswith((".docx", ".doc")):
            try:
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if not text.strip():
                    return "", f"Document appears to be empty: {filename}"
                return text, None
            except Exception as e:
                return "", f"Failed to parse Word document {filename}: {str(e)}"
                
        elif filename.lower().endswith(".pdf"):
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    
                    if len(pdf_reader.pages) == 0:
                        return "", f"PDF has no pages: {filename}"
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text += page_text + "\n"
                        except Exception as e:
                            print(f"couldnt extract page {page_num + 1} from {filename}: {e}")
                            continue
                    
                    if not text.strip():
                        return "", f"No text could be extracted from PDF: {filename}"
                        
                    return text, None
            except Exception as e:
                return "", f"Failed to parse PDF {filename}: {str(e)}"
        else:
            return "", f"Unsupported file format: {filename}"
            
    except Exception as e:
        return "", f"Unexpected error processing {filename}: {str(e)}"

# string matching algorithms implementation

class PerformanceMetrics:
    def __init__(self):
        self.comparisons = 0
        self.start_time = 0.0
        self.end_time = 0.0

    def start_timer(self): self.start_time = time.perf_counter()
    def stop_timer(self): self.end_time = time.perf_counter()
    @property
    def execution_time_ms(self) -> float: return (self.end_time - self.start_time) * 1000
    def reset(self):
        self.comparisons = 0
        self.start_time = 0.0
        self.end_time = 0.0

# brute force algorithm
def brute_force_search(text: str, pattern: str, metrics: PerformanceMetrics) -> bool:
    n, m = len(text), len(pattern)
    if m == 0: return True
    for i in range(n - m + 1):
        j = 0
        while j < m:
            metrics.comparisons += 1
            if text[i + j] != pattern[j]: break
            j += 1
        if j == m: return True
    return False

# rabin-karp algorithm with rolling hash
def rabin_karp_search(text: str, pattern: str, metrics: PerformanceMetrics) -> bool:
    n, m = len(text), len(pattern)
    if m == 0: return True
    if m > n: return False
    PRIME, D = 101, 256
    pattern_hash, text_hash, h = 0, 0, 1
    for i in range(m - 1): h = (h * D) % PRIME
    for i in range(m):
        pattern_hash = (D * pattern_hash + ord(pattern[i])) % PRIME
        text_hash = (D * text_hash + ord(text[i])) % PRIME
    for i in range(n - m + 1):
        if pattern_hash == text_hash:
            match = True
            for j in range(m):
                metrics.comparisons += 1
                if text[i + j] != pattern[j]:
                    match = False
                    break
            if match: return True
        if i < n - m:
            text_hash = (D * (text_hash - ord(text[i]) * h) + ord(text[i + m])) % PRIME
            if text_hash < 0: text_hash += PRIME
    return False

# kmp algorithm with lps preprocessing
def _compute_lps_array(pattern: str, m: int) -> List[int]:
    lps, length, i = [0] * m, 0, 1
    while i < m:
        if pattern[i] == pattern[length]:
            length += 1
            lps[i] = length
            i += 1
        else:
            if length != 0: length = lps[length - 1]
            else:
                lps[i] = 0
                i += 1
    return lps

def kmp_search(text: str, pattern: str, metrics: PerformanceMetrics) -> bool:
    n, m = len(text), len(pattern)
    if m == 0: return True
    lps, i, j = _compute_lps_array(pattern, m), 0, 0
    while i < n:
        metrics.comparisons += 1
        if pattern[j] == text[i]:
            i += 1
            j += 1
        if j == m: return True
        elif i < n and pattern[j] != text[i]:
            if j != 0: j = lps[j - 1]
            else: i += 1
    return False

# main analysis logic

def run_analysis(cv_text: str, keywords: List[str]) -> Dict[str, Any]:
    normalized_cv = normalize_text(cv_text)
    metrics = PerformanceMetrics()
    algorithms = [
        {"name": "Brute Force", "func": brute_force_search, "code": "BF"},
        {"name": "Rabin-Karp", "func": rabin_karp_search, "code": "RK"}, 
        {"name": "KMP", "func": kmp_search, "code": "KMP"},
    ]
    
    performance_results = []
    algorithm_matches = {}
    all_matched_keywords: Set[str] = set()
    
    # test each algorithm and collect results
    for algo in algorithms:
        metrics.reset()
        algo_matches = set()
        
        metrics.start_timer()
        for keyword in keywords:
            normalized_keyword = normalize_text(keyword)
            if algo["func"](normalized_cv, normalized_keyword, metrics):
                algo_matches.add(keyword)
        metrics.stop_timer()
        
        algorithm_matches[algo["code"]] = algo_matches
        all_matched_keywords.update(algo_matches)
        
        performance_results.append({
            "algorithm": algo["name"],
            "code": algo["code"],
            "execution_time_ms": round(metrics.execution_time_ms, 4),
            "comparisons": metrics.comparisons,
            "matches_found": len(algo_matches),
            "matched_keywords": sorted(list(algo_matches))
        })
    
    # calculate stats
    missing_keywords = set(keywords) - all_matched_keywords
    score = (len(all_matched_keywords) / len(keywords)) * 100 if keywords else 100
    
    # find best performing algorithms
    fastest_algo = min(performance_results, key=lambda x: x["execution_time_ms"])
    least_comparisons = min(performance_results, key=lambda x: x["comparisons"])
    
    return {
        "score": round(score, 2),
        "matched_count": len(all_matched_keywords),
        "total_keywords": len(keywords),
        "matched_keywords": sorted(list(all_matched_keywords)),
        "missing_keywords": sorted(list(missing_keywords)),
        "performance": performance_results,
        "algorithm_matches": {
            code: sorted(list(matches)) 
            for code, matches in algorithm_matches.items()
        },
        "performance_summary": {
            "fastest_algorithm": fastest_algo["algorithm"],
            "fastest_time_ms": fastest_algo["execution_time_ms"],
            "most_efficient_algorithm": least_comparisons["algorithm"],
            "least_comparisons": least_comparisons["comparisons"],
            "total_execution_time_ms": sum(p["execution_time_ms"] for p in performance_results),
            "total_comparisons": sum(p["comparisons"] for p in performance_results)
        }
    }


# api endpoints

@app.get("/", response_class=HTMLResponse)
async def get_frontend():
    logger.info("serving frontend page")
    try:
        with open("index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        logger.error("cant find index.html file")
        return HTMLResponse(content="<h1>Error: index.html not found.</h1>", status_code=500)

@app.post("/api/upload_cvs")
async def upload_cv_files(files: List[UploadFile] = File(...)):
    global cv_database
    
    success_count = 0
    failed_count = 0
    processed_files = []
    errors = []
    skipped_files = []

    # filter and validate files
    valid_files = []
    for file in files:
        filename = file.filename
        
        # check file type
        if not filename.lower().endswith((".docx", ".doc", ".pdf")):
            failed_count += 1
            errors.append(f"{filename}: Invalid file type (must be .pdf, .docx, or .doc)")
            continue
            
        # check filename format
        if not is_valid_cv_filename(filename):
            failed_count += 1
            errors.append(f"{filename}: Invalid filename format (expected: 2digits+letter+4digits, e.g., 23i2523)")
            continue
            
        valid_files.append(file)
    
    # get latest submissions only
    if valid_files:
        filenames = [f.filename for f in valid_files]
        latest_submissions = find_latest_submission(filenames)
        latest_filenames = set(latest_submissions.values())
        
        # process latest submissions
        for file in valid_files:
            filename = file.filename
            
            if filename not in latest_filenames:
                skipped_files.append(filename)
                continue
            
            text, error = await extract_text(file)
            
            if error:
                failed_count += 1
                errors.append(f"{filename}: {error}")
            elif not text.strip():
                failed_count += 1
                errors.append(f"{filename}: No text content found")
            else:
                cv_database[filename] = text
                processed_files.append(filename)
                success_count += 1
    
    message_parts = [f"{success_count} successful"]
    
    if failed_count > 0:
        message_parts.append(f"{failed_count} rejected")
    if skipped_files:
        message_parts.append(f"{len(skipped_files)} duplicates filtered")
        
    message = "Processing complete: " + ", ".join(message_parts) + "."
    
    response_data = {
        "message": message,
        "processed_files": processed_files,
        "errors": errors
    }
    
    if skipped_files:
        response_data["skipped_files"] = skipped_files
            
    return JSONResponse(content=response_data)

@app.get("/api/cvs")
async def get_cv_list():
    logger.info(f"getting cv list, have {len(cv_database)} cvs loaded")
    return JSONResponse(content=list(cv_database.keys()))

@app.get("/api/jds")
async def get_jd_list():
    logger.info(f"sending job descriptions, got {len(job_descriptions)} available")
    return JSONResponse(content=job_descriptions)

@app.get("/api/cv/{filename}")
async def get_cv_file(filename: str):
    try:
        # check if file exists in our database
        if filename not in cv_database:
            raise HTTPException(status_code=404, detail="CV not found")
        
        dataset_path = "./DataSet"
        file_path = os.path.join(dataset_path, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found on disk")
        
        # set content type
        if filename.lower().endswith('.pdf'):
            media_type = 'application/pdf'
        elif filename.lower().endswith(('.docx', '.doc')):
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            media_type = 'application/octet-stream'
        
        # read and send file
        with open(file_path, 'rb') as f:
            content = f.read()
        
        return Response(
            content=content,
            media_type=media_type,
            headers={"Content-Disposition": f"inline; filename={filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error serving file: {str(e)}")

@app.post("/api/load_dataset")
async def load_dataset_cvs():
    global cv_database
    logger.info("loading dataset files...")
    
    try:
        valid_files, filtered_files, rejected_files = load_dataset_files()
        
        if not filtered_files:
            return JSONResponse(content={
                "message": f"No valid CV files found in DataSet directory. {len(rejected_files)} files rejected due to naming convention.",
                "processed_files": [],
                "errors": ["DataSet directory is empty or contains no valid CV files"],
                "total_files_found": len(valid_files),
                "files_after_filtering": 0,
                "rejected_count": len(rejected_files)
            })
        
        success_count = 0
        failed_count = 0
        processed_files = []
        errors = []
        
        dataset_path = "./DataSet"
        
        for filename in filtered_files:
            file_path = os.path.join(dataset_path, filename)
            
            text, error = extract_text_from_file(file_path)
            
            if error:
                failed_count += 1
                errors.append(f"{filename}: {error}")
            elif not text.strip():
                failed_count += 1
                errors.append(f"{filename}: No text content found")
            elif len(text.strip()) < 50:
                failed_count += 1
                errors.append(f"{filename}: Text content too short ({len(text.strip())} chars)")
            else:
                # check if text has enough content
                words = text.strip().split()
                if len(words) < 10:
                    failed_count += 1
                    errors.append(f"{filename}: Text content insufficient ({len(words)} words)")
                else:
                    cv_database[filename] = text
                    processed_files.append(filename)
                    success_count += 1
        
        skipped_duplicates = len(valid_files) - len(filtered_files)
        rejected_count = len(rejected_files)
        
        message_parts = [f"{success_count} successful"]
        
        if failed_count > 0:
            message_parts.append(f"{failed_count} failed processing")
        if rejected_count > 0:
            message_parts.append(f"{rejected_count} rejected")
        if skipped_duplicates > 0:
            message_parts.append(f"{skipped_duplicates} duplicates filtered")
            
        message = "DataSet loaded: " + ", ".join(message_parts) + "."
        
        logger.info(f"dataset done: {success_count} ok, {failed_count} failed, {rejected_count} rejected, {skipped_duplicates} dupes")
        logger.info(f"total cvs loaded: {len(cv_database)}")
        
        return JSONResponse(content={
            "message": message,
            "processed_files": processed_files,
            "errors": errors,
            "total_files_found": len(valid_files) + len(rejected_files),
            "valid_files_found": len(valid_files),
            "files_after_filtering": len(filtered_files),
            "skipped_duplicates": skipped_duplicates,
            "rejected_count": rejected_count,
            "rejected_files": rejected_files[:10] if rejected_files else []
        })
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to load DataSet: {str(e)}"}
        )

class BatchAnalysisRequest(BaseModel):
    jd_id: str
    cv_filenames: List[str]

def calculate_dataset_analytics(results: List[Dict]) -> Dict[str, Any]:
    if not results:
        return {}
    
    # collect performance data for each algorithm
    algorithm_data = {
        "Brute Force": {"times": [], "comparisons": [], "matches": []},
        "Rabin-Karp": {"times": [], "comparisons": [], "matches": []},
        "KMP": {"times": [], "comparisons": [], "matches": []}
    }
    
    valid_results = [r for r in results if "full_report" in r and r["full_report"]]
    
    for result in valid_results:
        if "performance" in result["full_report"]:
            for perf in result["full_report"]["performance"]:
                algo_name = perf["algorithm"]
                if algo_name in algorithm_data:
                    algorithm_data[algo_name]["times"].append(perf["execution_time_ms"])
                    algorithm_data[algo_name]["comparisons"].append(perf["comparisons"])
                    algorithm_data[algo_name]["matches"].append(perf["matches_found"])
    
    # calculate stats for each algorithm
    analytics = {
        "total_cvs_analyzed": len(valid_results),
        "algorithms": {}
    }
    
    for algo_name, data in algorithm_data.items():
        if data["times"]:
            analytics["algorithms"][algo_name] = {
                "execution_time": {
                    "total_ms": round(sum(data["times"]), 4),
                    "average_ms": round(sum(data["times"]) / len(data["times"]), 4),
                    "min_ms": round(min(data["times"]), 4),
                    "max_ms": round(max(data["times"]), 4)
                },
                "comparisons": {
                    "total": sum(data["comparisons"]),
                    "average": round(sum(data["comparisons"]) / len(data["comparisons"]), 2),
                    "min": min(data["comparisons"]),
                    "max": max(data["comparisons"])
                },
                "matches_found": {
                    "total": sum(data["matches"]),
                    "average": round(sum(data["matches"]) / len(data["matches"]), 2),
                    "min": min(data["matches"]),
                    "max": max(data["matches"])
                },
                "files_processed": len(data["times"])
            }
    
    # find best algorithms
    if analytics["algorithms"]:
        fastest_algo = min(analytics["algorithms"].items(), 
                          key=lambda x: x[1]["execution_time"]["total_ms"])
        most_efficient = min(analytics["algorithms"].items(), 
                            key=lambda x: x[1]["comparisons"]["total"])
        
        analytics["summary"] = {
            "fastest_overall": fastest_algo[0],
            "fastest_total_time": fastest_algo[1]["execution_time"]["total_ms"],
            "most_efficient_overall": most_efficient[0],
            "least_total_comparisons": most_efficient[1]["comparisons"]["total"],
            "total_execution_time_all_algorithms": sum(
                data["execution_time"]["total_ms"] 
                for data in analytics["algorithms"].values()
            ),
            "total_comparisons_all_algorithms": sum(
                data["comparisons"]["total"] 
                for data in analytics["algorithms"].values()
            )
        }
    
    return analytics

@app.post("/api/analyze_batch")
async def analyze_batch(request: BatchAnalysisRequest):
    logger.info(f"starting analysis for jd {request.jd_id}, processing {len(request.cv_filenames)} cvs")
    
    jd = next((j for j in job_descriptions if j["id"] == request.jd_id), None)
    if not jd:
        logger.error(f"cant find job description: {request.jd_id}")
        raise HTTPException(status_code=404, detail="Job Description not found.")
        
    keywords = jd["keywords"]
    logger.info(f"using {jd['title']} with {len(keywords)} keywords")
    results = []
    
    for filename in request.cv_filenames:
        cv_text = cv_database.get(filename)
        if not cv_text:
            results.append({
                "cv_filename": filename,
                "score": 0,
                "error": "CV not found in database."
            })
            continue
            
        try:
            report = run_analysis(cv_text, keywords)
            
            summary = {
                "cv_filename": filename,
                "score": report["score"],
                "matched_count": report["matched_count"],
                "total_keywords": report["total_keywords"],
                "full_report": report  # Include full report for analytics
            }
            results.append(summary)
            
        except Exception as e:
            results.append({
                "cv_filename": filename,
                "score": 0,
                "error": f"Analysis rejected: {e}"
            })
            
    # sort by score (highest first)
    results.sort(key=lambda x: x.get("score", 0), reverse=True)
    
    # calculate overall analytics
    dataset_analytics = calculate_dataset_analytics(results)
    
    logger.info(f"analysis done, processed {len(results)} cvs")
    logger.info(f"calculated stats for {len(dataset_analytics.get('algorithms', {}))} algorithms")
    
    return JSONResponse(content={
        "results": results,
        "dataset_analytics": dataset_analytics,
        "job_description": {
            "title": jd["title"],
            "total_keywords": len(keywords),
            "keywords": keywords
        }
    })

# run the server
if __name__ == "__main__":
    print("CV Analyzer v2")
    print("starting server...")
    print("open http://127.0.0.1:8000 in browser")
    logger.info("server starting")
    logger.info(f"loaded {len(job_descriptions)} job descriptions")
    logger.info(f"cv database ready, {len(cv_database)} cvs loaded")
    uvicorn.run(app, host="127.0.0.1", port=8000)